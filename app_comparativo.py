import io
import json
from datetime import datetime
from flask import Flask, render_template_string, request, send_file
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)

# --- FRONTEND CON ALERTA VISUAL DE LÍMITES ---
INTERFACE_HTML = """
<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <title>Sistema de Tasación Inmobiliaria</title>
    <style>
        body { font-family: Arial, sans-serif; background-color: #f4f6f9; margin: 0; padding: 20px; color: #333; }
        .container { max-width: 1200px; margin: auto; background: white; padding: 20px; border-radius: 8px; box-shadow: 0 4px 10px rgba(0,0,0,0.1); }
        
        /* Barra de herramientas superior */
        .toolbar { background-color: #e9ecef; padding: 15px; border-radius: 6px; margin-bottom: 20px; display: flex; gap: 15px; align-items: center; }
        .btn { padding: 10px 18px; border: none; font-weight: bold; border-radius: 4px; cursor: pointer; color: white; font-size: 14px; }
        .btn-primary { background-color: #5b9bd5; } 
        .btn-success { background-color: #27ae60; } 
        .btn-danger { background-color: #e74c3c; }
        
        /* Bloque Objeto Consecuente (Fila 3 de tu Excel) */
        .bloque-superior { display: grid; grid-template-columns: repeat(auto-fit, minmax(180px, 1fr)); gap: 15px; margin-bottom: 20px; padding: 15px; border: 1px solid #ccc; border-radius: 4px; background-color: #fafafa; }
        .campo { display: flex; flex-direction: column; }
        .campo label { font-weight: bold; font-size: 12px; margin-bottom: 5px; color: #2c3e50; }
        .input-amarillo { background-color: #ffffcc; border: 1px solid #999; padding: 6px; text-align: center; font-weight: bold; }
        .input-verde { background-color: #e6ffe6; border: 1px solid #999; padding: 6px; text-align: center; font-weight: bold; color: green; }
        .input-azul { background-color: #cce6ff; border: 1px solid #999; padding: 6px; text-align: center; font-weight: bold; }
        .input-rosa { background-color: #ffe6e6; border: 1px solid #999; padding: 6px; text-align: center; font-weight: bold; color: #c0392b; }
        
        /* Tabla Principal */
        table { width: 100%; border-collapse: collapse; margin-top: 15px; font-size: 12px; }
        th, td { border: 1px solid #bdc3c7; padding: 6px; text-align: center; }
        th { background-color: #f0f0f0; font-weight: bold; color: #333; font-size: 11px; }
        .input-tabla { width: 90%; padding: 4px; border: 1px solid #ddd; text-align: center; border-radius: 3px; }
        
        /* Estilos de alerta para la celda fuera de rango */
        .celda-rango-ok { font-weight: bold; background-color: transparent; color: #333; transition: 0.3s; }
        .celda-rango-error { font-weight: bold; background-color: #ffcccc !important; color: #cc0000 !important; animation: parpadeo 1s infinite alternate; }
        
        @keyframes parpadeo {
            from { background-color: #ffcccc; }
            to { background-color: #ff9999; }
        }

        /* Diálogos modales */
        .modal { display: none; position: fixed; top: 0; left: 0; width: 100%; height: 100%; background: rgba(0,0,0,0.5); justify-content: center; align-items: center; }
        .modal-content { background: white; padding: 20px; border-radius: 6px; width: 350px; box-shadow: 0 4px 15px rgba(0,0,0,0.3); }
        .modal-content h3 { margin-top: 0; color: #2c3e50; border-bottom: 1px solid #eee; padding-bottom: 8px; }
        .fila-modal { margin-bottom: 12px; display: flex; justify-content: space-between; align-items: center; }
        .fila-modal input { width: 60%; padding: 5px; text-align: center; }
    </style>
</head>
<body>

<div class="container">
    <h2 style="text-align: center; color: #2c3e50; margin-top:0;">📊 TABLERO TÉCNICO DE VALUACIÓN METODO COMPARATIVO  (by ZWOL Soft)</h2>
    
    <div class="toolbar">
        <label><b>Cantidad de Antecedentes (1-10):</b></label>
        <input type="number" id="cant_antecedentes" value="3" min="1" max="10" style="width: 50px; padding: 5px; text-align: center; font-weight: bold;">
        <button class="btn btn-primary" onclick="configurarTabla()">INICIAR CONFIGURACIÓN</button>
        <button class="btn btn-danger" onclick="window.location.reload()">LIMPIAR TODO</button>
    </div>

    <form action="/generar-word" method="POST" id="form-tasacion">
        <div class="bloque-superior">
            <div class="campo">
                <label>OBJETO A TASAR (DIRECCIÓN):</label>
                <input type="text" name="direccion_objeto" value="Dirección, Localidad" style="padding:6px; font-weight:bold;">
            </div>
            <div class="campo">
                <label>TIPO DE INMUEBLE (R20):</label>
                <select name="tipo_inmueble" class="input-amarillo">
                    <option value="CASA" selected>CASA</option>
                    <option value="DEPARTAMENTO">DEPARTAMENTO</option>
                    <option value="LOCAL">LOCAL</option>
                    <option value="LOTE">LOTE</option>
                    <option value="QUINTA">QUINTA</option>
                    <option value="COCHERA">COCHERA</option>
                </select>
            </div>
            <div class="campo">
                <label>SUP. OBJETO m² (S3):</label>
                <input type="number" step="any" name="sup_objeto" id="sup_objeto" value="270" class="input-amarillo" oninput="calcularTodo()">
            </div>
            <div class="campo">
                <label>VALOR TOTAL USD (U3):</label>
                <input type="text" id="valor_total_venta" name="valor_total_venta" class="input-verde" readonly value="USD 0.00">
            </div>
            <div class="campo">
                <label>% RENT. AÑO (W3):</label>
                <input type="number" step="any" name="tasa_rentabilidad" id="tasa_rentabilidad" value="6.00" class="input-azul" oninput="calcularTodo()">
            </div>
            <div class="campo">
                <label>ALQ. MENSUAL SUGERIDO (Y3):</label>
                <input type="text" id="alquiler_sugerido" name="alquiler_sugerido" class="input-rosa" readonly value="USD 0.00">
            </div>
        </div>

        <div class="bloque-superior" style="display: flex; justify-content: space-between; gap: 10px; background-color: #f1f2f6; padding: 12px;">
            <div class="campo" style="flex: 1; min-width: 120px;">
                <label style="font-size: 11px; white-space: nowrap;">PROFESIONAL:</label>
                <input type="text" name="profesional" value="Nombre apellido" style="padding:4px; font-size:12px; width: 90%;">
            </div>
            <div class="campo" style="flex: 1; min-width: 100px;">
                <label style="font-size: 11px; white-space: nowrap;">VALOR DÓLAR ($):</label>
                <input type="number" step="any" name="valor_dolar" id="valor_dolar" value="1420" style="padding:4px; text-align:center; font-size:12px; width: 90%;" oninput="calcularTodo()">
            </div>
            <div class="campo" style="flex: 1; min-width: 110px;">
                <label style="font-size: 11px; white-space: nowrap;">ANTIGÜEDAD OBJETO:</label>
                <input type="number" name="antiguedad_obj" value="10" style="padding:4px; text-align:center; font-size:12px; width: 90%;">
            </div>
            <div class="campo" style="flex: 1; min-width: 110px;">
                <label style="font-size: 11px; white-space: nowrap;">FOS / FOT OBJETO:</label>
                <input type="text" name="fos_fot_obj" value="0.60 / 2.00" style="padding:4px; text-align:center; font-size:12px; width: 90%;">
            </div>
        </div>

        <table id="tabla-antecedentes">
            <thead>
                <tr>
                    <th>N°</th>
                    <th>UBICACIÓN</th>
                    <th>VALOR TOTAL (USD)</th>
                    <th>SUP. m²</th>
                    <th>VALOR m²</th>
                    <th>FOS</th>
                    <th>FOT</th>
                    <th>ACCIONES</th>
                    <th>COEF. TOTAL</th>
                    <th>VALOR HOMOG.</th>
                </tr>
            </thead>
            <tbody id="cuerpo-tabla"></tbody>
        </table>

        <input type="hidden" name="coeficientes_json" id="coeficientes_json">
        <input type="hidden" name="tabla_completa_json" id="tabla_completa_json">

        <br>
        <button type="submit" class="btn btn-success" style="width: 100%; padding: 14px; font-size: 16px;">🟩 GENERAR INFORME EN WORD COMPLETE (.DOCX)</button>
    </form>
</div>

<div class="modal" id="modal-coef">
    <div class="modal-content">
        <h3 id="modal-titulo">Coeficientes Ajuste</h3>
        <div id="campos-coeficientes"></div>
        <button class="btn btn-success" style="width:100%; margin-top:10px;" onclick="guardarCoeficientes()">Guardar Ajustes</button>
    </div>
</div>

<script>
let listaCoeficientes = {}; 
const nombresCoef = ["Actualización", "Ubicación", "Piso", "Planta", "Superficie", "Caract", "Edad", "Estado", "FF", "PT"];
let filaActualModificando = null;

function configurarTabla() {
    let cant = parseInt(document.getElementById('cant_antecedentes').value);
    let cuerpo = document.getElementById('cuerpo-tabla');
    cuerpo.innerHTML = "";
    
    for(let i=1; i<=cant; i++) {
        if (!listaCoeficientes[i]) {
            listaCoeficientes[i] = [1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0, 1.0];
        }
        
        let filaHtml = `
            <tr id="fila_${i}">
                <td><b>${i}</b></td>
                <td><input type="text" name="ubi_${i}" id="ubi_${i}" value="Antecedente Ejemplo ${i}" class="input-tabla" style="text-align:left;" oninput="calcularTodo()"></td>
                <td><input type="number" step="any" name="precio_${i}" id="precio_${i}" value="${100000 * i}" class="input-tabla" oninput="calcularFila(${i})"></td>
                <td><input type="number" step="any" name="sup_${i}" id="sup_${i}" value="${50 * i}" class="input-tabla" oninput="calcularFila(${i})"></td>
                <td><span id="val_m2_${i}">0.00</span></td>
                <td><input type="number" step="any" name="fos_${i}" id="fos_${i}" value="0.60" class="input-tabla" oninput="calcularTodo()"></td>
                <td><input type="number" step="any" name="fot_${i}" id="fot_${i}" value="3.00" class="input-tabla" oninput="calcularTodo()"></td>
                <td>
                    <button type="button" class="btn btn-primary" style="padding:4px 8px; font-size:10px;" onclick="abrirModalCoef(${i})">⚙️ AJUSTAR COEF (10)</button>
                </td>
                <td id="celda_coef_total_${i}" class="celda-rango-ok"><span id="coef_total_${i}">1.00</span></td>
                <td><span id="val_homog_${i}" style="font-weight:bold; color:blue;">0.00</span></td>
            </tr>
        `;
        cuerpo.insertAdjacentHTML('beforeend', filaHtml);
        calcularFila(i);
    }
}

function abrirModalCoef(id) {
    filaActualModificando = id;
    document.getElementById('modal-titulo').innerText = "Coeficientes - Antecedente " + id;
    let contenedor = document.getElementById('campos-coeficientes');
    contenedor.innerHTML = "";
    
    let coefs = listaCoeficientes[id];
    nombresCoef.forEach((nombre, index) => {
        let html = `
            <div class="fila-modal">
                <label>${nombre}:</label>
                <input type="number" step="0.01" id="modal_input_${index}" value="${coefs[index]}">
            </div>
        `;
        contenedor.insertAdjacentHTML('beforeend', html);
    });
    document.getElementById('modal-coef').style.display = 'flex';
}

function guardarCoeficientes() {
    let id = filaActualModificando;
    for(let index=0; index<10; index++) {
        listaCoeficientes[id][index] = parseFloat(document.getElementById(`modal_input_${index}`).value) || 1.0;
    }
    document.getElementById('modal-coef').style.display = 'none';
    calcularFila(id);
}

function calcularFila(id) {
    let precio = parseFloat(document.getElementById(`precio_${id}`).value) || 0;
    let sup = parseFloat(document.getElementById(`sup_${id}`).value) || 0;
    
    let valM2 = sup > 0 ? (precio / sup) : 0;
    document.getElementById(`val_m2_${id}`).innerText = valM2.toFixed(2);
    
    let coefs = listaCoeficientes[id];
    let coefTotal = coefs.reduce((a, b) => a * b, 1.0);
    
    document.getElementById(`coef_total_${id}`).innerText = coefTotal.toFixed(2);
    
    // --- SENSOR DE RANGO EXCEL (<0.7 o >1.3 se pinta de Rojo) ---
    let contenedorCelda = document.getElementById(`celda_coef_total_${id}`);
    if (coefTotal < 0.70 || coefTotal > 1.30) {
        contenedorCelda.className = "celda-rango-error";
    } else {
        contenedorCelda.className = "celda-rango-ok";
    }
    
    let valHomog = valM2 * coefTotal;
    document.getElementById(`val_homog_${id}`).innerText = valHomog.toFixed(2);
    
    calcularTodo();
}

function calcularTodo() {
    let cant = parseInt(document.getElementById('cant_antecedentes').value);
    let dataCompleta = [];
    let sumaHomog = 0;
    let contadorActivos = 0;
    
    for(let i=1; i<=cant; i++) {
        let elem = document.getElementById(`val_homog_${i}`);
        if(elem) {
            let vH = parseFloat(elem.innerText) || 0;
            sumaHomog += vH;
            contadorActivos++;
            
            dataCompleta.push({
                num: i,
                ubi: document.getElementById(`ubi_${i}`).value,
                precio: parseFloat(document.getElementById(`precio_${i}`).value) || 0,
                sup: parseFloat(document.getElementById(`sup_${i}`).value) || 0,
                val_m2: parseFloat(document.getElementById(`val_m2_${i}`).innerText) || 0,
                fos: document.getElementById(`fos_${i}`).value,
                fot: document.getElementById(`fot_${i}`).value,
                coef: parseFloat(document.getElementById(`coef_total_${i}`).innerText) || 1,
                homog: vH
            });
        }
    }
    
    let promedioHomog = contadorActivos > 0 ? (sumaHomog / contadorActivos) : 0;
    let supObjeto = parseFloat(document.getElementById('sup_objeto').value) || 0;
    
    let valorVentaFinal = promedioHomog * supObjeto;
    document.getElementById('valor_total_venta').value = "USD " + valorVentaFinal.toLocaleString('es-AR', {minimumFractionDigits: 2, maximumFractionDigits: 2});
    
    let rentabilidad = (parseFloat(document.getElementById('tasa_rentabilidad').value) || 0) / 100;
    let alqSugerido = (valorVentaFinal * rentabilidad) / 12;
    document.getElementById('alquiler_sugerido').value = "USD " + alqSugerido.toLocaleString('es-AR', {minimumFractionDigits: 2, maximumFractionDigits: 2});
    
    document.getElementById('coeficientes_json').value = JSON.stringify(listaCoeficientes);
    document.getElementById('tabla_completa_json').value = JSON.stringify(dataCompleta);
}

window.onload = function() {
    configurarTabla();
};
</script>
</body>
</html>
"""

@app.route('/')
def index():
    return render_template_string(INTERFACE_HTML)

@app.route('/generar-word', methods=['POST'])
def generar_word():
    direccion_objeto = request.form.get('direccion_objeto')
    tipo_inmueble = request.form.get('tipo_inmueble')
    sup_objeto = float(request.form.get('sup_objeto', 0))
    tasa_rent = float(request.form.get('tasa_rentabilidad', 0)) / 100
    valor_dolar = float(request.form.get('valor_dolar', 1))
    profesional = request.form.get('profesional')
    antiguedad_obj = request.form.get('antiguedad_obj')
    fos_fot_obj = request.form.get('fos_fot_obj')
    
    coeficientes_completos = json.loads(request.form.get('coeficientes_json', '{}'))
    nombres_coef = ["Actualización", "Ubicación", "Piso", "Planta", "Superficie", "Caract.", "Edad", "Estado", "F/F", "P/T"]

    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Encabezado
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_titulo.add_run("INFORME TÉCNICO DE VALUACIÓN INMOBILIARIA")
    run_t.font.name = 'Arial'; run_t.font.size = Pt(14); run_t.bold = True
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(f"PROFESIONAL: {profesional.upper()} | FECHA: {datetime.now().strftime('%d/%m/%Y')}")
    run_sub.font.name = 'Arial'; run_sub.font.size = Pt(10)

    # 1. Características Objeto
    h1 = doc.add_paragraph()
    r_h1 = h1.add_run("1. CARACTERÍSTICAS DEL INMUEBLE OBJETO")
    r_h1.font.name = 'Arial'; r_h1.font.size = Pt(12); r_h1.bold = True
    
    p_obj = doc.add_paragraph()
    p_obj.add_run(f"Dirección: {direccion_objeto} ({tipo_inmueble})\n").bold = True
    p_obj.add_run(f"• Superficie Total: {sup_objeto} m²\n")
    p_obj.add_run(f"• Antigüedad: {antiguedad_obj} Años\n")
    p_obj.add_run(f"• F.O.S. / F.O.T.: {fos_fot_obj}\n")

    # 2. Memoria Técnica
    h2 = doc.add_paragraph()
    r_h2 = h2.add_run("2. MEMORIA TÉCNICA Y FORMULEO")
    r_h2.font.name = 'Arial'; r_h2.font.size = Pt(12); r_h2.bold = True
    
    p_mem = doc.add_paragraph()
    p_mem.add_run("Se utiliza el Método Comparativo de Mercado, aplicando la siguiente expresión de homogeneización para cada antecedente:\n\n")
    
    p_form = doc.add_paragraph()
    p_form.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_form = p_form.add_run("V.H. = V.B. x (C1 x C2 x C3 x ... x Cn)")
    r_form.font.italic = True; r_form.bold = True; r_form.font.size = Pt(11)
    
    doc.add_paragraph("Donde V.H. es el Valor Homogeneizado, V.B. el Valor Base y los 'Cn' son los coeficientes de ajuste aplicados por producto (10 indicadores).\n\n")

    suma_homog = 0
    contador_ant = 0
    
    for key, coefs in coeficientes_completos.items():
        ubi = request.form.get(f'ubi_{key}')
        precio = float(request.form.get(f'precio_{key}', 0))
        sup = float(request.form.get(f'sup_{key}', 0))
        fos_ant = request.form.get(f'fos_{key}', '0.60')
        fot_ant = request.form.get(f'fot_{key}', '3.00')
        
        if not ubi or precio <= 0 or sup <= 0:
            continue
            
        contador_ant += 1
        valor_base_m2 = precio / sup
        
        coef_total = 1.0
        for c in coefs:
            coef_total *= c
            
        valor_homog_m2 = valor_base_m2 * coef_total
        suma_homog += valor_homog_m2
        
        p_ant = doc.add_paragraph()
        p_ant.add_run(f"ANTECEDENTE N°{key} ({ubi})\n").bold = True
        
        p_ind = doc.add_paragraph()
        r_ind = p_ind.add_run(f"Indicadores del Antecedente: F.O.S.: {fos_ant} | F.O.T.: {fot_ant}\n")
        r_ind.font.italic = True; r_ind.font.size = Pt(10)
        
        texto_coefs = "Coeficientes aplicados:\n"
        for idx, nombre in enumerate(nombres_coef):
            texto_coefs += f"  • {nombre}: {coefs[idx]:.2f}\n"
        texto_coefs += f"  • COE TOTAL: {coef_total:.2f}\n"
        p_ant.add_run(texto_coefs)
        
        p_calc_A = doc.add_paragraph()
        run_A = p_calc_A.add_run(f"A) Cálculo Valor Base: USD {precio:,.2f} / {sup} m² = USD {valor_base_m2:.4f} / m²\n")
        run_A.font.name = 'Arial'; run_A.font.size = Pt(10); run_A.font.italic = True
        
        cadena_coef = " x ".join([f"{c:.2f}" for c in coefs])
        p_calc_B = doc.add_paragraph()
        run_B = p_calc_B.add_run(f"B) Aplicación: ({valor_base_m2:.4f}) x ({cadena_coef}) = USD {valor_homog_m2:,.2f} / m²\n")
        run_B.font.name = 'Arial'; run_B.font.size = Pt(10); run_B.font.italic = True

    # 3. Dictamen Final
    promedio_homog = suma_homog / contador_ant if contador_ant > 0 else 0
    valor_venta_final = promedio_homog * sup_objeto
    alquiler_mensual_usd = (valor_venta_final * tasa_rent) / 12
    
    h3 = doc.add_paragraph()
    r_h3 = h3.add_run("3. DICTAMEN FINAL Y VALORES")
    r_h3.font.name = 'Arial'; r_h3.font.size = Pt(12); r_h3.bold = True
    
    doc.add_paragraph("El valor final surge del promedio simple de los valores resultantes:")
    
    p_prom = doc.add_paragraph()
    p_prom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_prom = p_prom.add_run(f"Cálculo: USD {suma_homog:,.2f} / {contador_ant} = USD {promedio_homog:,.2f} / m²")
    run_prom.font.name = 'Arial'; run_prom.bold = True
    
    p_res_v = doc.add_paragraph()
    p_res_v.add_run("\nVALOR DE VENTA SUGERIDO:\n").bold = True
    p_res_v.add_run(f"• TOTAL USD: USD {valor_venta_final:,.2f}\n").bold = True
    p_res_v.add_run(f"• En Pesos: $ {valor_venta_final * valor_dolar:,.2f}\n")
    
    p_res_a = doc.add_paragraph()
    p_res_a.add_run("ALQUILER MENSUAL ESTIMADO:\n").bold = True
    p_res_a.add_run(f"• TOTAL USD: USD {alquiler_mensual_usd:,.2f}\n").bold = True
    p_res_a.add_run(f"• En Pesos: $ {alquiler_mensual_usd * valor_dolar:,.2f}\n\n")

    target = io.BytesIO()
    doc.save(target)
    target.seek(0)
    
    return send_file(
        target,
        as_attachment=True,
        download_name=f"Informe_Tasacion_{direccion_objeto.replace(' ', '_')}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=10000, debug=True)
